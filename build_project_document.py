from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path("AI_PDF_Summarizer_Project_Documentation.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_FILL = "F2F4F7"
BORDER = "A6A6A6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), BORDER)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def paragraph_border_bottom(paragraph, color="2E74B5", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_para(doc, text="", size=11, bold=False, italic=False, color=None, align=None, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color is not None:
            r.font.color.rgb = color
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(16 if level == 1 else 13)
    run.bold = True
    run.font.color.rgb = BLUE if level == 1 else DARK_BLUE
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(item)
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(item)
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.style = doc.styles["No Spacing"]
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instr)
    run._r.append(field_end)


def add_index_table(doc):
    data = [
        ("SERIAL NO", "TOPIC", "PAGE NO"),
        ("1", "INTRODUCTION", "3"),
        ("2", "METHODOLOGY", "4-5"),
        ("3", "FINDING AND ANALYSIS", "6"),
        ("4", "CODE AND OUTCOMES", "7-10"),
        ("5", "KEY TAKEAWAYS", "11"),
        ("6", "CONCLUSION", "12"),
    ]
    table = doc.add_table(rows=len(data), cols=3)
    table.autofit = False
    set_table_width(table, [1.0, 4.2, 1.3])
    set_cell_margins(table)
    set_table_borders(table)
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx]
        for c_idx, value in enumerate(row_data):
            set_cell_text(row.cells[c_idx], value, bold=(r_idx == 0), size=10.5)
            row.cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx == 0:
                set_cell_shading(row.cells[c_idx], LIGHT_FILL)


def add_info_table(doc, rows, widths=(1.8, 4.7)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    set_table_width(table, list(widths))
    set_cell_margins(table)
    set_table_borders(table)
    for idx, (label, value) in enumerate(rows):
        set_cell_text(table.rows[idx].cells[0], label, bold=True)
        set_cell_text(table.rows[idx].cells[1], value)
        set_cell_shading(table.rows[idx].cells[0], LIGHT_FILL)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    # Cover page
    add_para(doc, "COLLEGE NAME & LOGO", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    add_para(doc, "PROJECT NOTEBOOK", size=24, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=26)
    p = add_para(doc, "Project Title :-", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    paragraph_border_bottom(p)
    add_para(doc, "AI PDF SUMMARIZER", size=20, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    add_para(doc, "Submitted by-", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_bullets(
        doc,
        [
            "Anurag Kumar",
            "Avinash Kumar",
            "Amlan Roy",
            "Aishwarya Mukherjee",
            "Aishanee Das",
            "Anushka Dutta",
            "Ankit Kumar",
        ],
    )
    add_para(doc, "", after=6)
    add_info_table(
        doc,
        [
            ("REGISTRATION NO", "____________________________"),
            ("UNIVERSITY ROLL NO", "____________________________"),
            ("SEMESTER", "____________________________"),
            ("STREAM", "____________________________"),
            ("YEAR OF SUBMISSION", "2026"),
        ],
    )
    doc.add_page_break()

    add_para(doc, "INDEX", size=20, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_index_table(doc)
    doc.add_page_break()

    add_heading(doc, "INTRODUCTION: -", 1)
    add_para(
        doc,
        "In recent years, Artificial Intelligence has become an important tool for reading, searching, and understanding large documents. Students, teachers, and professionals often need to study long PDF files such as notes, reports, research papers, and manuals. Reading every page manually can take a lot of time, especially when the user only needs the main points or answers to specific questions.",
    )
    add_para(
        doc,
        "Our project, AI PDF Summarizer, is a Streamlit-based application that allows users to upload a PDF file, extract page-wise text, generate AI-powered summaries, ask questions about selected pages, and view useful document analytics. The application uses Python for processing, PyMuPDF for PDF extraction, Pandas and Matplotlib for analysis and charts, and Google Gemini API for summary generation and question answering.",
    )
    add_para(
        doc,
        "The main objective of this project is to reduce the effort required to understand lengthy PDF documents. It provides a simple web interface where users can inspect extracted text, select page ranges, choose the summary style, and download the result in different formats.",
    )
    doc.add_page_break()

    add_heading(doc, "METHODOLOGY: -", 1)
    add_heading(doc, "1. Problem Definition", 2)
    add_para(
        doc,
        "The project aims to develop an AI-assisted tool that can read PDF documents and produce structured summaries. The problem addressed is the difficulty of quickly understanding long documents without manually reading every page.",
    )
    add_heading(doc, "2. Source of Data", 2)
    add_para(
        doc,
        "The source data for this project is the PDF file uploaded by the user. The PDF may contain study notes, reports, project documents, articles, or other text-based material. The application extracts text page by page and uses that extracted content for analytics and AI processing.",
    )
    add_heading(doc, "3. PDF Text Extraction", 2)
    add_bullets(
        doc,
        [
            "The uploaded PDF is read using PyMuPDF.",
            "Each page is processed separately to collect page number, text, word count, and character count.",
            "The extracted text is cleaned by removing extra spaces and unnecessary line breaks.",
            "If selected pages contain no readable text, the application informs the user that OCR is required.",
        ],
    )
    add_heading(doc, "4. AI Summary Generation", 2)
    add_para(
        doc,
        "The cleaned PDF text is divided into manageable chunks so that longer documents can be processed reliably. Each chunk is sent to the Gemini model with a prompt based on the selected summary style. The partial summaries are then combined into one final structured summary.",
    )
    add_heading(doc, "5. User Interface and Export", 2)
    add_para(
        doc,
        "The user interface is created using Streamlit. It includes options for selecting page ranges, choosing summary style, changing chunk size, viewing analytics, asking questions, and exporting summaries as TXT, Markdown, or JSON.",
    )
    doc.add_page_break()

    add_heading(doc, "TOOLS & TECHNOLOGIES: -", 1)
    add_info_table(
        doc,
        [
            ("Programming Language", "Python"),
            ("Frontend / Interface", "Streamlit"),
            ("PDF Processing", "PyMuPDF / fitz"),
            ("Data Analysis", "Pandas"),
            ("Visualization", "Matplotlib"),
            ("AI Model", "Google Gemini API"),
            ("Configuration", "python-dotenv"),
            ("Export Formats", "TXT, Markdown, JSON, CSV"),
        ],
    )
    add_heading(doc, "DATA CLEANING: -", 1)
    add_bullets(
        doc,
        [
            "Extra spaces, blank lines, and irregular text breaks are cleaned from extracted PDF text.",
            "Empty pages are counted separately so the user can identify pages without readable content.",
            "Page ranges entered by the user are validated and filtered against the total number of pages.",
            "Word count and character count are calculated for every page.",
            "Keyword frequency is calculated after removing common stopwords.",
        ],
    )
    doc.add_page_break()

    add_heading(doc, "FINDING AND ANALYSIS: -", 1)
    add_heading(doc, "1. Understanding the Document", 2)
    add_para(
        doc,
        "The application works with uploaded PDF documents and converts them into page-wise text data. This structure helps the application show statistics and create page-aware summaries with references.",
    )
    add_heading(doc, "2. Analytics Generated", 2)
    add_bullets(
        doc,
        [
            "Total number of selected pages.",
            "Total word count across selected pages.",
            "Estimated reading time based on word count.",
            "Number of empty pages.",
            "Top keywords from the document text.",
            "Bar chart showing words per page.",
        ],
    )
    add_heading(doc, "3. Key Observations", 2)
    add_bullets(
        doc,
        [
            "Page-wise extraction makes it easier to summarize only the required pages.",
            "Word count chart helps users identify dense and important pages.",
            "Keyword frequency provides a quick view of the document's main topics.",
            "AI-based question answering helps users find answers without searching manually.",
        ],
    )
    add_heading(doc, "4. Practical Implications", 2)
    add_para(
        doc,
        "The project can be useful for students preparing notes, teachers reviewing documents, and users who need fast understanding of large reports. It can also be extended with OCR for scanned PDFs and better document search features.",
    )
    doc.add_page_break()

    add_heading(doc, "CODE AND OUTCOMES: -", 1)
    add_heading(doc, "INPUT: - Main Application Imports", 2)
    add_code_block(
        doc,
        [
            "import matplotlib.pyplot as plt",
            "import pandas as pd",
            "import streamlit as st",
            "",
            "from src.analytics import calculate_document_stats, keyword_frequency",
            "from src.pdf_reader import extract_pdf_pages, filter_pages, parse_page_ranges",
            "from src.summarizer import GeminiSummarizer",
        ],
    )
    add_heading(doc, "PDF Text Extraction Code", 2)
    add_code_block(
        doc,
        [
            "with fitz.open(stream=pdf_bytes, filetype=\"pdf\") as pdf:",
            "    for page_number, page in enumerate(pdf, start=1):",
            "        text = page.get_text(\"text\", sort=True).strip()",
            "        pages.append({",
            "            \"page\": page_number,",
            "            \"text\": text,",
            "            \"word_count\": len(text.split()),",
            "            \"char_count\": len(text),",
            "        })",
        ],
    )
    add_para(
        doc,
        "OUTPUT 1: The program reads the uploaded PDF and creates a page-wise dataset containing page number, extracted text, word count, and character count.",
        bold=True,
    )
    doc.add_page_break()

    add_heading(doc, "Analytics and Chart Code", 2)
    add_code_block(
        doc,
        [
            "fig, ax = plt.subplots(figsize=(10, 4))",
            "ax.bar(df[\"page\"], df[\"word_count\"], color=\"#2f6f73\")",
            "ax.set_xlabel(\"Page Number\")",
            "ax.set_ylabel(\"Word Count\")",
            "ax.set_title(\"Words Per Page\")",
            "st.pyplot(fig)",
        ],
    )
    add_para(
        doc,
        "OUTPUT GRAPH 1: Bar chart showing the number of words on each selected PDF page.",
        bold=True,
    )
    add_heading(doc, "AI Summary Code", 2)
    add_code_block(
        doc,
        [
            "summary = summarizer.summarize_pdf(",
            "    pages,",
            "    style=summary_style,",
            "    max_chars=max_chars,",
            "    progress_callback=progress.progress,",
            ")",
        ],
    )
    add_para(
        doc,
        "OUTPUT 2: The Gemini model generates a structured summary based on the selected summary style such as Academic, Brief, Detailed, Executive, Study Notes, or Explain Simply.",
        bold=True,
    )
    doc.add_page_break()

    add_heading(doc, "Ask PDF Code", 2)
    add_code_block(
        doc,
        [
            "answer = summarizer.answer_question(pages, question, max_chars=max_chars)",
            "st.subheader(\"Answer\")",
            "st.write(answer)",
        ],
    )
    add_para(
        doc,
        "OUTPUT 3: The user can ask a question and receive an answer based only on the extracted PDF context.",
        bold=True,
    )
    add_heading(doc, "Export Code", 2)
    add_code_block(
        doc,
        [
            "col1.download_button(label=\"Download TXT\", data=summary)",
            "col2.download_button(label=\"Download Markdown\", data=summary_to_markdown(summary))",
            "col3.download_button(label=\"Download JSON\", data=summary_to_json(summary, source_file))",
        ],
    )
    add_para(
        doc,
        "OUTPUT 4: The generated summary can be downloaded in TXT, Markdown, or JSON format. Page statistics can also be exported as CSV.",
        bold=True,
    )
    doc.add_page_break()

    add_heading(doc, "PROJECT FEATURES: -", 1)
    add_bullets(
        doc,
        [
            "Upload any text-based PDF document.",
            "Select specific pages or page ranges for analysis.",
            "Generate summaries in different writing styles.",
            "Ask questions about the uploaded PDF.",
            "View total pages, total words, reading time, and empty page count.",
            "Display word-count chart and top keywords.",
            "Download summaries and page statistics.",
        ],
    )
    add_heading(doc, "EXPECTED SCREEN OUTPUTS: -", 1)
    add_bullets(
        doc,
        [
            "PDF Statistics panel with selected pages, total words, reading time, and empty pages.",
            "Summary tab showing the AI generated summary.",
            "Ask PDF tab showing answers to user questions.",
            "Analytics tab showing words per page and keyword frequency.",
            "Extracted Text tab showing the cleaned text preview.",
        ],
    )
    doc.add_page_break()

    add_heading(doc, "KEY TAKEAWAYS: -", 1)
    add_bullets(
        doc,
        [
            "Gained practical experience in building a Python-based AI web application.",
            "Learned how to extract and clean text from PDF documents.",
            "Understood how to connect a Streamlit application with the Gemini API.",
            "Implemented document analytics using Pandas and Matplotlib.",
            "Created useful export features for TXT, Markdown, JSON, and CSV formats.",
            "Improved teamwork and collaboration by dividing work across interface, logic, documentation, and testing.",
        ],
    )
    add_heading(doc, "CHALLENGES FACED: -", 1)
    add_bullets(
        doc,
        [
            "Some PDFs may contain scanned images instead of selectable text, which requires OCR.",
            "Long PDFs need chunking so the AI model can process the content properly.",
            "Page range input must be validated to avoid invalid page selections.",
            "API key configuration must be handled carefully through the .env file.",
        ],
    )
    doc.add_page_break()

    add_heading(doc, "CONCLUSION", 1)
    add_para(
        doc,
        "The AI PDF Summarizer project provided a valuable opportunity to apply Python, data analysis, visualization, and artificial intelligence to a practical document-processing problem. The project demonstrates how a user can upload a PDF, extract its text, understand document statistics, generate summaries, ask questions, and export useful outputs from one simple application.",
    )
    add_para(
        doc,
        "Through this project, our team learned how to combine different technologies such as Streamlit, PyMuPDF, Pandas, Matplotlib, and Google Gemini API. The application is useful for students and professionals who need to save time while reading large documents.",
    )
    add_para(
        doc,
        "Overall, this project improved our understanding of AI-powered applications, PDF processing, prompt-based summarization, and user-friendly software development. In the future, the project can be improved by adding OCR support for scanned PDFs, user login, document history, and more advanced search features.",
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
